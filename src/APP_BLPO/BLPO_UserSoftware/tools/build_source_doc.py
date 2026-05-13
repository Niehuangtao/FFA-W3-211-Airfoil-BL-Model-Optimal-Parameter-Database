from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
OUT = DOCS / "sample源程序.docx"
DOC_OUT = DOCS / "sample源程序.doc"
TEXT_OUT = DOCS / "source_program_payload.txt"


SOURCE_FILES = [
    ROOT / "BLPO_AppLauncher.m",
    ROOT / "BLPO_App.m",
    ROOT / "BLPO_sample_config.m",
    ROOT / "BLPO_make_sample_data.m",
    ROOT / "BLPO_selftest.m",
    ROOT / "run_BLPO_cli.m",
    ROOT / "package_BLPO_app.m",
]
SOURCE_FILES.extend(sorted((ROOT / "functions").glob("*.m"), key=lambda p: p.name.lower()))


def set_run_font(run, name="Courier New", east_asia="宋体", size=8, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold


def add_page_number(section):
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    paragraph._p.append(fld_begin)
    paragraph._p.append(instr)
    paragraph._p.append(fld_end)


def read_source(path):
    for encoding in ("utf-8", "gbk"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def add_centered(doc, text, size=12, bold=False, space_after=6):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(space_after)
    run = paragraph.add_run(text)
    set_run_font(run, name="Arial", east_asia="宋体", size=size, bold=bold)
    return paragraph


def add_code_line(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(8.6)
    run = paragraph.add_run(text)
    set_run_font(run, size=7.5)


def make_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.header_distance = Inches(0.2)
    add_page_number(section)

    add_centered(doc, "软件全称    Beddoes–Leishman动态失速模型参数优化软件", size=12)
    add_centered(doc, "软件简称    BLPO", size=12)
    add_centered(doc, "软件版本    V1.0", size=12)
    add_centered(doc, "", size=12, space_after=18)
    add_centered(doc, "源 程 序", size=18, bold=True, space_after=24)
    add_centered(doc, "上海交通大学", size=12)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_centered(doc, "源 程 序 代 码 部 分", size=14, bold=True, space_after=12)

    total_lines = 0
    for path in SOURCE_FILES:
        if not path.exists():
            continue
        rel = path.relative_to(ROOT).as_posix()
        add_code_line(doc, f"%% ===== File: {rel} =====")
        for raw_line in read_source(path).splitlines():
            line = raw_line.rstrip()
            if not line.strip():
                continue
            add_code_line(doc, line)
            total_lines += 1

    add_code_line(doc, f"%% ===== End of source code, nonblank source lines: {total_lines} =====")
    doc.save(OUT)
    source_text = make_source_text()
    TEXT_OUT.write_text(source_text, encoding="utf-8")
    DOC_OUT.write_text(make_rtf(source_text), encoding="ascii")
    print(OUT)


def clean_text(text):
    cleaned = []
    for ch in text:
        code = ord(ch)
        if ch in "\t\r\n" or code >= 32:
            cleaned.append(ch)
        else:
            cleaned.append(" ")
    return "".join(cleaned)


def make_source_text():
    lines = [
        "软件全称    Beddoes–Leishman动态失速模型参数优化软件",
        "软件简称    BLPO",
        "软件版本    V1.0",
        "",
        "源 程 序",
        "",
        "上海交通大学",
        "",
        "(源 程 序 代 码 部 分)",
    ]
    total_lines = 0
    for path in SOURCE_FILES:
        if not path.exists():
            continue
        rel = path.relative_to(ROOT).as_posix()
        lines.append(f"%% ===== File: {rel} =====")
        for raw_line in clean_text(read_source(path)).splitlines():
            line = raw_line.rstrip()
            if not line.strip():
                continue
            lines.append(line)
            total_lines += 1
    lines.append(f"%% ===== End of source code, nonblank source lines: {total_lines} =====")
    return "\n".join(lines)


def rtf_escape(text):
    out = []
    for ch in text:
        if ch == "\\":
            out.append(r"\\")
        elif ch == "{":
            out.append(r"\{")
        elif ch == "}":
            out.append(r"\}")
        elif ch == "\t":
            out.append(r"\tab ")
        elif ch == "\n":
            out.append(r"\par" + "\n")
        else:
            code = ord(ch)
            if 32 <= code <= 126:
                out.append(ch)
            elif code >= 32:
                signed = code if code < 32768 else code - 65536
                out.append(rf"\u{signed}?")
            else:
                out.append(" ")
    return "".join(out)


def make_rtf(source_text):
    lines = source_text.splitlines()
    rtf = [
        r"{\rtf1\ansi\deff0",
        r"{\fonttbl{\f0 Courier New;}{\f1 SimSun;}}",
        r"\paperw11906\paperh16838\margl720\margr720\margt540\margb540",
        r"\fs15",
    ]
    for line in lines:
        if line in {"源 程 序"}:
            rtf.append(r"\pard\qc\f1\fs36\b " + rtf_escape(line) + r"\b0\par")
        elif line.startswith("软件") or line == "上海交通大学":
            rtf.append(r"\pard\qc\f1\fs24 " + rtf_escape(line) + r"\par")
        else:
            rtf.append(r"\pard\f0\fs15 " + rtf_escape(line) + r"\par")
    rtf.append("}")
    return "\n".join(rtf)


if __name__ == "__main__":
    make_doc()
