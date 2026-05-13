from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "assets"
OUT = DOCS / "用户手册-BLPO-交互软件版.docx"


def font(run, size=None, bold=False):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if size:
        run.font.size = Pt(size)
    run.bold = bold


def make_doc():
    doc = Document()
    for name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        if name == "Normal":
            style.font.size = Pt(11)

    def p(text, style=None, align=None, size=None, bold=False):
        para = doc.add_paragraph(style=style) if style else doc.add_paragraph()
        run = para.add_run(text)
        font(run, size=size, bold=bold)
        if align is not None:
            para.alignment = align
        para.paragraph_format.space_after = Pt(6)
        return para

    def bullet(text):
        para = doc.add_paragraph(style="List Bullet")
        run = para.add_run(text)
        font(run)
        para.paragraph_format.space_after = Pt(2)

    def num(text):
        para = doc.add_paragraph(style="List Number")
        run = para.add_run(text)
        font(run)
        para.paragraph_format.space_after = Pt(2)

    def image(name, caption):
        p(caption, align=WD_ALIGN_PARAGRAPH.CENTER)
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run().add_picture(str(ASSETS / name), width=Inches(6.4))
        para.paragraph_format.space_after = Pt(10)

    p("Beddoes-Leishman动态失速模型参数优化软件", align=WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True)
    p("软件简称：BLPO    软件版本：V1.0", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    p("用 户 手 册", "Title", WD_ALIGN_PARAGRAPH.CENTER)
    p("上海交通大学", align=WD_ALIGN_PARAGRAPH.CENTER)

    p("目 录", "Heading 1")
    for line in [
        "一、概况",
        "二、系统简介",
        "三、软件安装说明",
        "四、软件操作流程",
        "1 软件启动",
        "2 软件操作过程及界面",
        "（1）主界面组成",
        "（2）参考数据设置",
        "（3）翼型与工况设置",
        "（4）优化参数设置",
        "（5）默认模型预览",
        "（6）运行参数优化",
        "（7）导出优化结果",
        "3 软件小结",
    ]:
        p(line)

    p("一、概况", "Heading 1")
    p("1 编写目的", "Heading 2")
    p("本软件用于翼型动态失速模型参数标定。用户通过界面输入翼型运动工况、选择参考数据文件，即可调用 Beddoes-Leishman 动态失速模型并开展参数优化。软件解决了传统脚本流程中需要反复修改工况、文件路径和优化参数的问题，提高了参数标定的效率和可重复性。")
    p("软件支持两种参考数据：试验数据为两列格式（time、Cl），CFD 数据为四列格式（time、Cl、Cd、Cm）。软件根据列数自动识别数据类型，并以最后一个周期回滞环的非重叠面积作为优化目标，同时保留 RMSE 作为辅助评价指标。")
    p("2 软件、硬件运行环境", "Heading 2")
    bullet("操作系统：Windows 10 或 Windows 11。")
    bullet("运行环境：MATLAB R2025a，或安装 MATLAB Runtime 后运行打包程序。")
    bullet("优化工具箱：Global Optimization Toolbox；并行加速可选 Parallel Computing Toolbox。")
    p("3 编程语言版本号及程序容量", "Heading 2")
    bullet("编程语言：MATLAB。")
    bullet("开发版本：MATLAB R2025a。")
    bullet("程序容量：约 3000 行 MATLAB 代码，包含 GUI、数据读取、BL 模型计算、优化求解、结果导出和文档生成脚本。")
    p("4 软件功能与特点", "Heading 2")
    bullet("可由用户自定义平均攻角、俯仰幅值、缩减频率、来流速度、雷诺数、弦长和计算步数。")
    bullet("可导入试验数据或 CFD 数据作为优化参考。")
    bullet("可预览默认 BL 模型结果，便于检查数据范围和初始拟合情况。")
    bullet("可按 full、lag、cl、cm 四种模式执行遗传算法参数优化，其中 full 模式按 lag、cl、cm 三个阶段顺序优化。")
    bullet("可导出参数表、误差指标、逐点对比数据、MAT 结果文件和拟合曲线图。")
    p("5 报价", "Heading 2")
    p("10000元（人民币）。")

    p("二、系统简介", "Heading 1")
    p("1 概述", "Heading 2")
    p("BLPO 动态失速模型参数优化软件是一套面向翼型非定常气动分析和工程应用的交互式工具。软件围绕 BL 模型参数标定流程构建，将工况输入、数据导入、模型计算、回滞环非重叠面积评价、参数优化和结果导出集成在同一界面内。")
    p("2 应用范围和对象", "Heading 2")
    p("软件适用于风力机翼型、潮流能水轮机翼型、直升机旋翼翼型及其他存在动态失速问题的翼型，可用于试验数据拟合、CFD 结果对比和工程气动模型参数修正。")
    p("3 系统特色", "Heading 2")
    bullet("图形化输入工况，减少源程序修改。")
    bullet("自动识别 2 列或 4 列参考数据。")
    bullet("优化结果与参考数据同屏对比。")
    bullet("运行过程和导出文件清晰，便于软著登记和工程归档。")

    p("三、软件安装说明", "Heading 1")
    p("1 安装产品", "Heading 2")
    num("安装 MATLAB R2025a。")
    num("确认已安装 Global Optimization Toolbox。")
    num("如需运行独立程序，安装与 R2025a 匹配的 MATLAB Runtime。")
    num("将 BLPO_UserSoftware 文件夹放置到本地磁盘，并确保 output 文件夹可写。")
    p("2 启动前准备", "Heading 2")
    num("准备参考数据文件。试验数据为 time、Cl 两列；CFD 数据为 time、Cl、Cd、Cm 四列。")
    num("准备翼型静态数据文件，或使用 sample_data 中的默认翼型文件。")
    num("确认输出目录存在，默认输出目录为 BLPO_UserSoftware/output。")

    p("四、软件操作流程", "Heading 1")
    p("1 软件启动", "Heading 2")
    p("打开 MATLAB，将当前路径切换到 BLPO_UserSoftware 工作目录，在命令窗口输入 BLPO_AppLauncher 并回车，即可启动图形界面。若使用打包程序，可直接双击 dist 文件夹下的 BLPO.exe。")
    num("启动 MATLAB R2025a。")
    num("进入 BLPO_UserSoftware 工作目录。")
    num("输入 BLPO_AppLauncher。")
    num("等待主界面显示。")

    p("2 软件操作过程及界面", "Heading 2")
    p("进入主程序后，界面左侧为参数输入区，右侧为结果显示区。左侧包含工况设置、参考数据设置、输出目录设置和优化设置；右侧以 2×2 形式显示 Cl 时间历程、Cl 回滞环、Cm 时间历程和 Cm 回滞环，并在右侧保留参数表和运行日志。界面左下方设置进度条，预览、优化和导出过程中会同步显示当前进度与状态。")
    image("screenshot_main.png", "图1 BLPO 软件主界面")
    image("flowchart_operation.png", "图2 BLPO 软件总体使用流程")

    p("（1）主界面组成", "Heading 3")
    p("左侧 Case setup 面板分为三个部分。Motion condition 用于输入运动工况；Reference and output 用于选择参考数据、翼型静态数据和输出目录；Optimization 用于设置优化模式、遗传算法种群规模、最大迭代代数和局部优化选项。右侧绘图区包括四幅图：左上为最后三个周期的 Cl 时间历程，右上为最后一个周期的 Cl 闭合回滞环，左下为最后三个周期的 Cm 时间历程，右下为最后一个周期的 Cm 闭合回滞环。Parameter 表格显示当前参数值。")

    p("（2）参考数据设置", "Heading 3")
    p("在 Reference file 输入框中设置参考数据文件。用户可点击右侧“...”按钮选择文件，也可直接粘贴文件路径。软件读取文件后会自动判断列数：若为 2 列，则采用 Cl 回滞环非重叠面积作为优化目标；若为 4 列，则 lag 和 cl 阶段采用 Cl 面积目标，cm 阶段采用 Cm 面积目标。")
    num("点击 Reference file 右侧按钮。")
    num("选择试验数据或 CFD 数据文件。")
    num("确认路径显示在输入框中。")
    num("检查数据第一列是否为时间，且时间范围与模拟工况有重叠。")
    p("参考数据应尽量只保留数值列。若文件含有复杂表头、单位说明或非数值字符，建议先整理为 csv、txt、dat 或 out 格式。")

    p("（3）翼型与工况设置", "Heading 3")
    p("在 Airfoil polar 输入框中设置翼型静态数据文件。默认文件为 sample_data/FFA_W3_211.txt。工况参数在 Motion condition 中设置，其中 Mean angle 表示平均攻角，Amplitude 表示俯仰幅值，Phase 表示相位，Reynolds number 表示雷诺数，Velocity U 表示来流速度，Reduced freq. k 表示缩减频率，Chord c 表示弦长，Cycles 表示模拟周期数，Steps per cycle 表示每周期计算步数。")
    p("工况设置完成后，建议先点击 Preview default model，检查参考数据与模型输出是否出现在同一时间范围内。若曲线为空或软件提示时间范围不重叠，应检查参考数据时间单位、Cycles 和 Steps per cycle 设置。")

    p("（4）优化参数设置", "Heading 3")
    p("Optimization 区域用于控制优化过程。Mode 下拉菜单包含 full、lag、cl、cm 四种模式。选择 lag、cl 或 cm 时，软件只优化对应类别的参数，其余参数保持默认值；选择 full 时，软件先优化 lag 参数，再以上一步结果为初值优化 cl 参数，最后继续优化 cm 参数。Population size 为遗传算法种群规模，Max generations 为每个优化阶段的最大迭代代数。")
    p("首次检查流程时，建议使用较小的种群规模和迭代代数，例如 Population size 设为 8 或 16，Max generations 设为 1 至 5。正式优化时可根据计算资源提高参数。Run local search after GA 勾选后会在遗传算法结束后继续执行局部优化，Use parallel pool when available 勾选后可调用并行计算资源。")
    image("flowchart_algorithm.png", "图3 BL 参数优化计算流程")

    p("（5）默认模型预览", "Heading 3")
    p("点击 Preview default model 后，软件会根据当前输入构造攻角时间序列，读取参考数据，调用默认 BL 参数计算 Cl、Cd、Cm，并在右侧 2×2 绘图区显示 Cl 时间历程、Cl 回滞环、Cm 时间历程和 Cm 回滞环。时间历程图显示最后三个周期，回滞环显示最后一个周期并闭合。预览用于确认数据读取是否正确、模拟时间是否覆盖参考数据、默认参数与参考数据差异是否合理。")
    num("确认 Reference file、Airfoil polar 和 Motion condition 已设置完成。")
    num("点击 Preview default model。")
    num("观察日志区是否出现 Preview complete。")
    num("查看右侧 Cl 和 Cm 的时间历程图及闭合回滞环。")
    num("查看 Area objective、Cl 面积和 Cm 面积指标。")
    image("screenshot_preview.png", "图4 默认模型预览结果")

    p("（6）运行参数优化", "Heading 3")
    p("点击 Run optimization 后，软件按照当前模式调用遗传算法。若选择 full，进度条会依次显示 lag、cl、cm 三个阶段，每个阶段均使用 Max generations 设置的代数；若选择 lag、cl 或 cm，则只执行对应参数组的一次优化。优化目标为参考回滞环与模型回滞环的非重叠面积，且每个阶段只有在面积目标不增大时才接受新参数。优化过程中，MATLAB 命令窗口会显示每一代的最优目标函数值和平均目标函数值；界面左下角进度条会按照初始化、遗传算法迭代、局部优化、结果绘制和文件导出的顺序更新；日志区会显示优化开始和结束状态。优化完成后，软件会用最优参数重新计算模型响应，并在曲线图中同时显示参考数据、初始参数计算结果和优化参数计算结果，参数表也会同时列出初始值和当前优化值。")
    num("选择优化模式。快速测试可选择 lag 或 cl，完整计算可选择 full。")
    num("设置 Population size 和 Max generations。")
    num("根据需要选择是否执行局部优化和并行计算。")
    num("点击 Run optimization。")
    num("等待 Status 显示 optimization complete。")
    num("查看参数表和曲线变化，判断优化结果是否可接受。")
    image("screenshot_optimization.png", "图5 参数优化完成界面")

    p("（7）导出优化结果", "Heading 3")
    p("优化完成后，软件会自动导出结果。用户也可点击 Export last result，将最近一次预览或优化结果重新保存到 Output folder 指定目录。默认输出目录为 BLPO_UserSoftware/output。")
    bullet("optimized_parameters.csv：保存参数名称、初始参数值和最优参数值。")
    bullet("fit_metrics.csv：保存 Cl 非重叠面积、Cm 非重叠面积、面积目标函数以及 Cl、Cd、Cm 的 RMSE 辅助指标。")
    bullet("comparison_data.csv：保存参考数据、初始参数计算结果和优化参数计算结果的逐点对比。")
    bullet("BLPO_result.mat：保存完整 MATLAB 结果结构体。")
    bullet("fit_comparison.png：保存 Cl 和 Cm 的时间历程及回滞环对比图。")
    p("若需要将优化参数用于其他程序，可直接读取 optimized_parameters.csv；若需要继续在 MATLAB 中分析，可读取 BLPO_result.mat。")

    p("3 软件小结", "Heading 2")
    p("BLPO 动态失速模型参数优化软件将工况输入、参考数据导入、模型计算、参数优化、结果分析和文件导出整合为一个完整的交互流程。用户无需修改源程序即可完成常规优化任务，既便于工程应用，也便于对计算过程进行归档和复核。")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    make_doc()
